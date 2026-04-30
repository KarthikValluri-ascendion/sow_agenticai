"""Apply PII redactions to .docx using python-docx; handles run fragmentation."""

from __future__ import annotations

import io
from typing import BinaryIO

from docx import Document
from docx.oxml.ns import qn


def _replace_paragraph_text(paragraph, mapping: dict[str, str]) -> None:
    """
    Replace keys in paragraph.text with values, then rewrite runs as a single run.
    Word often splits text across multiple w:r elements; paragraph.text is the logical merge.
    """
    if not mapping:
        return
    text = paragraph.text
    if not text:
        return
    new_text = text
    for old, new in sorted(mapping.items(), key=lambda x: -len(x[0])):
        if old:
            new_text = new_text.replace(old, new)
    if new_text == text:
        return
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:r"):
            p.remove(child)
    paragraph.add_run(new_text)


def _walk_cell(cell, fn) -> None:
    for p in cell.paragraphs:
        fn(p)
    for table in cell.tables:
        for row in table.rows:
            for c in row.cells:
                _walk_cell(c, fn)


def _apply_to_document_body(doc: Document, mapping: dict[str, str]) -> None:
    def fn(p):
        _replace_paragraph_text(p, mapping)

    for p in doc.paragraphs:
        fn(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _walk_cell(cell, fn)


def _apply_to_header_footer(doc: Document, mapping: dict[str, str]) -> None:
    def fn(p):
        _replace_paragraph_text(p, mapping)

    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                fn(p)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _walk_cell(cell, fn)


def redact_docx_bytes(source: BinaryIO, mapping: dict[str, str]) -> io.BytesIO:
    """
    Load a .docx from a binary stream, apply replacements, return BytesIO buffer.
    Iterates main body (paragraphs + tables, nested tables) and header/footer regions.
    """
    doc = Document(source)
    _apply_to_document_body(doc, mapping)
    _apply_to_header_footer(doc, mapping)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
